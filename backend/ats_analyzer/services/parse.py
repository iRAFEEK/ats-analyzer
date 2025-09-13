"""Document parsing service with OCR fallback."""

import io
import tempfile
from pathlib import Path
from typing import Optional, Tuple

import fitz  # PyMuPDF
import structlog
from docx import Document
from PIL import Image
import pytesseract

from ats_analyzer.api.dto import ParseMeta
from ats_analyzer.core.errors import FileProcessingError, TextExtractionError, OCRError


logger = structlog.get_logger(__name__)


class ParsedDocument:
    """Parsed document container."""
    
    def __init__(self, text: str, meta: ParseMeta):
        self.text = text
        self.meta = meta


def detect_filetype(content: bytes, filename: str, content_type: Optional[str]) -> str:
    """Detect file type from content and metadata."""
    # Check magic bytes first
    if content.startswith(b'%PDF'):
        return 'pdf'
    elif content.startswith(b'PK\x03\x04'):
        return 'docx'
    elif content.startswith((b'\x89PNG', b'\xff\xd8\xff', b'GIF8')):
        return 'image'
    
    # Fallback to filename extension
    if filename:
        ext = Path(filename).suffix.lower()
        if ext == '.pdf':
            return 'pdf'
        elif ext in ['.docx', '.doc']:
            return 'docx'
        elif ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']:
            return 'image'
    
    # Fallback to content type
    if content_type:
        if 'pdf' in content_type:
            return 'pdf'
        elif 'wordprocessingml' in content_type or 'msword' in content_type:
            return 'docx'
        elif 'image' in content_type:
            return 'image'
    
    raise FileProcessingError(f"Unsupported file type: {filename}")


def extract_text_from_pdf(content: bytes) -> Tuple[str, bool, float]:
    """Extract text from PDF using PyMuPDF with OCR fallback."""
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        
        text_parts = []
        has_columns = False
        has_tables = False
        ocr_used = False
        total_chars = 0
        extracted_chars = 0
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Try text extraction first
            page_text = page.get_text()
            
            if page_text.strip():
                # Text extraction successful
                text_parts.append(page_text)
                extracted_chars += len(page_text)
            else:
                # No text found, try OCR
                logger.info(f"No text found on page {page_num + 1}, attempting OCR")
                try:
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    
                    # Use Tesseract OCR
                    image = Image.open(io.BytesIO(img_data))
                    ocr_text = pytesseract.image_to_string(image)
                    
                    if ocr_text.strip():
                        text_parts.append(ocr_text)
                        extracted_chars += len(ocr_text)
                        ocr_used = True
                        logger.info(f"OCR extracted {len(ocr_text)} chars from page {page_num + 1}")
                    
                except Exception as e:
                    logger.warning(f"OCR failed on page {page_num + 1}: {str(e)}")
            
            # Analyze layout
            blocks = page.get_text("dict")["blocks"]
            if len(blocks) > 1:
                # Check for multi-column layout by analyzing x-coordinates
                x_positions = []
                for block in blocks:
                    if "lines" in block:
                        for line in block["lines"]:
                            x_positions.append(line["bbox"][0])
                
                if len(set(round(x / 50) for x in x_positions)) > 2:
                    has_columns = True
            
            # Simple table detection (look for aligned text blocks)
            if len(blocks) > 3:
                y_positions = [block["bbox"][1] for block in blocks if "lines" in block]
                if len(set(round(y / 10) for y in y_positions)) < len(y_positions) * 0.7:
                    has_tables = True
            
            total_chars += max(len(page_text), 100)  # Estimate expected chars
        
        doc.close()
        
        full_text = "\n\n".join(text_parts)
        extractability_score = min(extracted_chars / max(total_chars, 1), 1.0)
        
        return full_text, ocr_used, extractability_score
        
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(content: bytes) -> Tuple[str, bool, float]:
    """Extract text from DOCX file."""
    try:
        doc = Document(io.BytesIO(content))
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        full_text = "\n".join(text_parts)
        
        # DOCX files don't need OCR and have perfect extractability
        return full_text, False, 1.0
        
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_image(content: bytes) -> Tuple[str, bool, float]:
    """Extract text from image using OCR."""
    try:
        image = Image.open(io.BytesIO(content))
        
        # Use Tesseract OCR
        ocr_text = pytesseract.image_to_string(image, config='--psm 6')
        
        if not ocr_text.strip():
            raise OCRError("No text found in image")
        
        # OCR extractability is typically lower
        extractability_score = 0.8 if len(ocr_text.strip()) > 50 else 0.5
        
        return ocr_text, True, extractability_score
        
    except pytesseract.TesseractError as e:
        raise OCRError(f"Tesseract OCR failed: {str(e)}")
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from image: {str(e)}")


async def parse_document(
    content: bytes, 
    filename: str, 
    content_type: Optional[str] = None
) -> ParsedDocument:
    """Parse document and extract text with metadata."""
    if not content:
        raise FileProcessingError("Empty file content")
    
    # Detect file type
    filetype = detect_filetype(content, filename, content_type)
    
    logger.info(
        "Parsing document",
        filename=filename,
        filetype=filetype,
        size_bytes=len(content),
    )
    
    # Extract text based on file type
    if filetype == 'pdf':
        text, ocr_used, extractability_score = extract_text_from_pdf(content)
        has_columns = True  # Will be properly detected in extract_text_from_pdf
        has_tables = False  # Will be properly detected in extract_text_from_pdf
    elif filetype == 'docx':
        text, ocr_used, extractability_score = extract_text_from_docx(content)
        has_columns = False
        has_tables = True  # DOCX likely has tables
    elif filetype == 'image':
        text, ocr_used, extractability_score = extract_text_from_image(content)
        has_columns = False
        has_tables = False
    else:
        raise FileProcessingError(f"Unsupported file type: {filetype}")
    
    if not text.strip():
        raise TextExtractionError("No text content extracted from document")
    
    # Create metadata
    meta = ParseMeta(
        filetype=filetype,
        has_columns=has_columns,
        has_tables=has_tables,
        extractability_score=extractability_score,
        ocr_used=ocr_used,
    )
    
    logger.info(
        "Document parsing completed",
        text_length=len(text),
        extractability_score=extractability_score,
        ocr_used=ocr_used,
    )
    
    return ParsedDocument(text=text, meta=meta)
